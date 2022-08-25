import os
import sys
import pandas as pd
from docx import Document


def docwrite(data, name):
    doc = Document()
    if 'year' in data.columns and 'birth' in data.columns and 'loc' in data.columns:
        cat = 'lj'
    elif 'birthdate' in data.columns and 'location' in data.columns and 'year' in data.columns:
        cat = 'vk'
    for i, row in data.iterrows():
        if str(row['left']) != 'nan':
            p = doc.add_paragraph(f"{i}\t{row['left']}")
        else:
            p = doc.add_paragraph(f'{i}\t')
        p.add_run(f" {row['result']} ").bold = True
        if str(row['right']) != 'nan':
            p.add_run(str(row['right']))
        if cat == 'lj':
            doc.add_paragraph(f"Дата создания текста: {row['year']}")
            doc.add_paragraph(f"Год рождения автора: {row['birth']}")
            doc.add_paragraph(f"Место жительства автора: {row['loc']}")
        elif cat == 'vk':
            doc.add_paragraph(f"Дата создания текста: {row['year']}")
            doc.add_paragraph(f"Год рождения автора: {row['birthdate']}")
            doc.add_paragraph(f"Место жительства автора: {row['location']}")       
        doc.add_paragraph(' ')
    doc.save(f'docs/{name}.docx')

if not os.path.exists('docs'):
    os.mkdir('docs')

if not os.path.exists('data'):
    os.mkdir('data')
    print('Put your xslx or txt with tab-separated snippets into data folder')
    sys.exit()

files = os.listdir('data')

for name in files:
    if name.endswith('.xlsx'):
        data = pd.read_excel(f'data/{name}')
        data = data.drop_duplicates(subset='left', keep='first')
        docwrite(data, name[:-5])
    elif name.endswith('.txt'):
        data = pd.read_csv(f'data/{name}', delimiter='\t', header=0, index_col='int_id', quoting=3, on_bad_lines='warn')
        data = data.drop_duplicates(subset='left', keep='first')
        docwrite(data, name[:-4])